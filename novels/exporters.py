"""
Document Export Pipeline — Phase 7.5
Generates .docx and .epub files from approved chapters.
"""

import os
import logging
from pathlib import Path
from datetime import datetime
from django.conf import settings

logger = logging.getLogger(__name__)

EXPORTS_DIR = Path(settings.BASE_DIR) / 'exports'


class BookExporter:
    """
    Exports a book's approved chapters to .docx and .epub formats.
    Injects KDP metadata from KeywordResearch + BookDescription.
    """

    def __init__(self, book):
        self.book = book
        self.output_dir = EXPORTS_DIR / str(book.id)
        self.output_dir.mkdir(parents=True, exist_ok=True)

        # Lazy-load related data
        self._keyword_data = None
        self._description = None
        self._chapters = None

    @property
    def keyword_data(self):
        if self._keyword_data is None:
            from novels.models import KeywordResearch
            self._keyword_data = KeywordResearch.objects.filter(book=self.book).first()
        return self._keyword_data

    @property
    def active_description(self):
        if self._description is None:
            from novels.models import BookDescription
            self._description = (
                BookDescription.objects.filter(book=self.book, is_active=True).first()
                or BookDescription.objects.filter(book=self.book).first()
            )
        return self._description

    @property
    def approved_chapters(self):
        if self._chapters is None:
            from novels.models import Chapter, ChapterStatus
            self._chapters = list(
                Chapter.objects.filter(book=self.book, status=ChapterStatus.APPROVED)
                .order_by('chapter_number')
            )
        return self._chapters

    def get_metadata(self) -> dict:
        """Assemble all metadata for injection into the document."""
        kw = self.keyword_data
        pen = self.book.pen_name

        title = (kw.suggested_title if kw else None) or self.book.title
        subtitle = kw.suggested_subtitle if kw else None
        author = pen.name if pen else 'Unknown Author'
        keywords = kw.kdp_backend_keywords if kw else []
        description_plain = ''
        if self.active_description:
            # Strip HTML tags for EPUB metadata
            import re
            description_plain = re.sub(r'<[^>]+>', '', self.active_description.description_html)

        return {
            'title': title,
            'subtitle': subtitle,
            'author': author,
            'keywords': keywords,
            'description': description_plain,
            'publisher': pen.name if pen else 'AI Novel Factory',
            'language': 'en',
            'date': datetime.now().strftime('%Y-%m-%d'),
        }

    def get_legal_disclaimer(self) -> str:
        """Generate the standard legal disclaimer."""
        pen = self.book.pen_name
        author_name = pen.name if pen else 'the Author'
        year = datetime.now().year

        return (
            f"Copyright © {year} {author_name}. All rights reserved.\n\n"
            "This is a work of fiction. Names, characters, places, and incidents are either the "
            "product of the author's imagination or are used fictitiously. Any resemblance to "
            "actual persons, living or dead, events, or locales is entirely coincidental.\n\n"
            "No part of this publication may be reproduced, stored in a retrieval system, or "
            "transmitted in any form or by any means—electronic, mechanical, photocopying, "
            "recording, or otherwise—without the prior written permission of the publisher.\n\n"
            "This book contains content that was written with AI assistance."
        )

    # -------------------------------------------------------------------------
    # DOCX Export
    # -------------------------------------------------------------------------

    def export_docx(self) -> str:
        """
        Export approved chapters to a formatted .docx file.
        Returns path to the generated file.
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError("python-docx is required. Run: pip install python-docx")

        meta = self.get_metadata()
        doc = Document()

        # --- Document styles ---
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)

        # --- Title page ---
        title_para = doc.add_heading(meta['title'], level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if meta.get('subtitle'):
            sub_para = doc.add_paragraph(meta['subtitle'])
            sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub_para.runs[0].bold = True

        author_para = doc.add_paragraph(f"by {meta['author']}")
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_page_break()

        # --- Legal disclaimer ---
        doc.add_heading('Copyright', level=2)
        doc.add_paragraph(self.get_legal_disclaimer())
        doc.add_page_break()

        # --- Chapters ---
        chapters = self.approved_chapters
        if not chapters:
            logger.warning(f"No approved chapters for book {self.book.id}")

        for chapter in chapters:
            doc.add_heading(f'Chapter {chapter.chapter_number}', level=1)
            if chapter.title:
                sub = doc.add_heading(chapter.title, level=2)

            # Add chapter content (split into paragraphs)
            content = chapter.content or ''
            for para_text in content.split('\n\n'):
                para_text = para_text.strip()
                if para_text:
                    doc.add_paragraph(para_text)

            doc.add_page_break()

        # --- Save ---
        safe_title = "".join(c for c in meta['title'] if c.isalnum() or c in ' _-')[:50]
        filename = f"{safe_title}.docx"
        filepath = self.output_dir / filename
        doc.save(str(filepath))

        logger.info(f"DOCX exported to {filepath}")
        return str(filepath)

    # -------------------------------------------------------------------------
    # EPUB Export
    # -------------------------------------------------------------------------

    def export_epub(self) -> str:
        """
        Export approved chapters to EPUB format.
        Returns path to the generated file.
        """
        try:
            from ebooklib import epub
        except ImportError:
            raise ImportError("ebooklib is required. Run: pip install ebooklib")

        meta = self.get_metadata()
        book_epub = epub.EpubBook()

        # --- Metadata ---
        book_epub.set_title(meta['title'])
        book_epub.set_language(meta['language'])
        book_epub.add_author(meta['author'])
        book_epub.add_metadata('DC', 'publisher', meta['publisher'])
        book_epub.add_metadata('DC', 'date', meta['date'])
        if meta['description']:
            book_epub.add_metadata('DC', 'description', meta['description'])
        for kw in meta['keywords']:
            book_epub.add_metadata('DC', 'subject', kw)

        # --- CSS Style ---
        css_content = """
body { font-family: Georgia, serif; font-size: 1em; line-height: 1.6; margin: 5%; }
h1 { font-size: 1.8em; margin-top: 2em; }
h2 { font-size: 1.4em; margin-top: 1.5em; }
p { text-indent: 1.5em; margin: 0.3em 0; }
.chapter-title { page-break-before: always; }
.disclaimer { font-size: 0.85em; font-style: italic; color: #555; }
"""
        css = epub.EpubItem(
            uid='style_default',
            file_name='style/default.css',
            media_type='text/css',
            content=css_content.encode('utf-8'),
        )
        book_epub.add_item(css)

        spine = ['nav']
        toc = []

        # --- Legal page ---
        disclaimer_html = f"""
<html><head><link rel="stylesheet" href="../style/default.css"/></head><body>
<div class="disclaimer">
<h2>Copyright Notice</h2>
<pre style="white-space:pre-wrap;">{self.get_legal_disclaimer()}</pre>
</div>
</body></html>
"""
        disclaimer_ch = epub.EpubHtml(
            title='Copyright',
            file_name='copyright.xhtml',
            lang='en',
        )
        disclaimer_ch.content = disclaimer_html.encode('utf-8')
        disclaimer_ch.add_item(css)
        book_epub.add_item(disclaimer_ch)
        spine.append(disclaimer_ch)

        # --- Chapters ---
        for chapter in self.approved_chapters:
            content = chapter.content or ''
            # Convert paragraphs to <p> tags
            paragraphs_html = '\n'.join(
                f'<p>{p.strip()}</p>'
                for p in content.split('\n\n')
                if p.strip()
            )

            chapter_html = f"""
<html><head><link rel="stylesheet" href="../style/default.css"/></head><body>
<div class="chapter-title"><h1>Chapter {chapter.chapter_number}</h1></div>
{f'<h2>{chapter.title}</h2>' if chapter.title else ''}
{paragraphs_html}
</body></html>
"""
            epub_ch = epub.EpubHtml(
                title=f'Chapter {chapter.chapter_number}',
                file_name=f'chapter_{chapter.chapter_number:03d}.xhtml',
                lang='en',
            )
            epub_ch.content = chapter_html.encode('utf-8')
            epub_ch.add_item(css)
            book_epub.add_item(epub_ch)
            spine.append(epub_ch)
            toc.append(epub.Link(
                f'chapter_{chapter.chapter_number:03d}.xhtml',
                f'Chapter {chapter.chapter_number}',
                f'chapter{chapter.chapter_number}',
            ))

        # Nav and spine
        book_epub.toc = toc
        book_epub.add_item(epub.EpubNcx())
        book_epub.add_item(epub.EpubNav())
        book_epub.spine = spine

        # --- Save ---
        safe_title = "".join(c for c in meta['title'] if c.isalnum() or c in ' _-')[:50]
        filename = f"{safe_title}.epub"
        filepath = self.output_dir / filename
        epub.write_epub(str(filepath), book_epub, {})

        logger.info(f"EPUB exported to {filepath}")
        return str(filepath)

    def export_summary(self) -> dict:
        """Return a summary of what will be exported."""
        chapters = self.approved_chapters
        total_words = sum(len((ch.content or '').split()) for ch in chapters)
        return {
            'title': self.book.title,
            'author': self.book.pen_name.name if self.book.pen_name else 'Unknown',
            'approved_chapters': len(chapters),
            'total_words': total_words,
            'output_dir': str(self.output_dir),
        }
